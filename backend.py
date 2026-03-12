"""
VAPT Report Generator — Ethical Byte
Flask backend with DOCX + PDF generation
"""

import os, io, json, sqlite3, hashlib, secrets, base64, re
from datetime import datetime, timedelta
from functools import wraps
from flask import Flask, request, jsonify, send_file, redirect, render_template_string, send_from_directory, session

# ── ReportLab ──────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
pt = 1.0  # ReportLab works natively in points
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, Image, HRFlowable
)
from reportlab.platypus.flowables import Flowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.graphics.shapes import Drawing, Rect, String, Line
from reportlab.graphics import renderPDF

# ── python-docx ────────────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Pillow ─────────────────────────────────────────────────────────────────
from PIL import Image as PILImage

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

# ───────────────────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder='.')
app.secret_key = os.environ.get('SECRET_KEY', 'vapt-secret-change-me-in-production')

BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
DB_PATH    = os.path.join(BASE_DIR, 'reports.db')
REPORT_DIR = os.path.join(BASE_DIR, 'generated_reports')
LOGO_PATH  = os.path.join(BASE_DIR, 'logo.png')
TEMPLATE_PATH = os.path.join(BASE_DIR, 'Report_Template.docx')
os.makedirs(REPORT_DIR, exist_ok=True)

SUPERUSER_EMAIL = os.environ.get('SUPERUSER_EMAIL', 'admin@ethicalbyte.com')
SUPERUSER_PASS  = os.environ.get('SUPERUSER_PASS',  'Admin@1234')
OTP_EXPIRY_MIN  = int(os.environ.get('OTP_EXPIRY_MIN', 10))

SEV_COLORS = {
    'Critical': '#C00000',
    'High':     '#FF4444',
    'Medium':   '#FFC000',
    'Low':      '#92D050',
}

# ══════════════════════════════════════════════════════════════════════════
#  DATABASE
# ══════════════════════════════════════════════════════════════════════════

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as db:
        db.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            mobile TEXT,
            is_superuser INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            org_name TEXT,
            format TEXT,
            filename TEXT,
            filepath TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            FOREIGN KEY(user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS otps (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            mobile TEXT NOT NULL,
            otp TEXT NOT NULL,
            name TEXT, email TEXT, password_hash TEXT,
            expires_at TEXT NOT NULL
        );
        """)
        # Ensure superuser exists
        exists = db.execute('SELECT id FROM users WHERE email=?', (SUPERUSER_EMAIL,)).fetchone()
        if not exists:
            db.execute(
                'INSERT INTO users(name,email,password_hash,is_superuser) VALUES(?,?,?,1)',
                ('Admin', SUPERUSER_EMAIL, hashlib.sha256(SUPERUSER_PASS.encode()).hexdigest())
            )

init_db()

# ══════════════════════════════════════════════════════════════════════════
#  AUTH HELPERS
# ══════════════════════════════════════════════════════════════════════════

def hash_pw(pw): return hashlib.sha256(pw.encode()).hexdigest()

def login_required(f):
    @wraps(f)
    def wrapped(*a, **kw):
        if 'user_id' not in session:
            if request.is_json:
                return jsonify({'error': 'Not authenticated'}), 401
            return redirect('/login')
        return f(*a, **kw)
    return wrapped

def superuser_required(f):
    @wraps(f)
    def wrapped(*a, **kw):
        if 'user_id' not in session:
            return jsonify({'error': 'Not authenticated'}), 401
        with get_db() as db:
            u = db.execute('SELECT is_superuser FROM users WHERE id=?', (session['user_id'],)).fetchone()
        if not u or not u['is_superuser']:
            return jsonify({'error': 'Forbidden'}), 403
        return f(*a, **kw)
    return wrapped

# ══════════════════════════════════════════════════════════════════════════
#  STATIC HTML ROUTES
# ══════════════════════════════════════════════════════════════════════════

@app.route('/login')
def login_page():
    return send_from_directory(BASE_DIR, 'login.html')

@app.route('/')
@login_required
def index():
    return send_from_directory(BASE_DIR, 'frontend.html')

@app.route('/admin')
@superuser_required
def admin_page():
    return send_from_directory(BASE_DIR, 'admin.html')

@app.route('/logo.png')
def logo():
    return send_from_directory(BASE_DIR, 'logo.png')

# ══════════════════════════════════════════════════════════════════════════
#  AUTH API
# ══════════════════════════════════════════════════════════════════════════

@app.route('/api/auth/send-otp', methods=['POST'])
def send_otp():
    data = request.json or {}
    mobile   = data.get('mobile','').strip()
    name     = data.get('name','').strip()
    email    = data.get('email','').strip().lower()
    password = data.get('password','')
    if not all([mobile, name, email, password]):
        return jsonify({'error': 'All fields required'}), 400
    otp = str(secrets.randbelow(900000) + 100000)
    expires = (datetime.utcnow() + timedelta(minutes=OTP_EXPIRY_MIN)).isoformat()
    pw_hash = hash_pw(password)
    with get_db() as db:
        db.execute('DELETE FROM otps WHERE mobile=? OR email=?', (mobile, email))
        db.execute(
            'INSERT INTO otps(mobile,otp,name,email,password_hash,expires_at) VALUES(?,?,?,?,?,?)',
            (mobile, otp, name, email, pw_hash, expires)
        )
    # Try Twilio
    sid   = os.environ.get('TWILIO_ACCOUNT_SID','')
    token = os.environ.get('TWILIO_AUTH_TOKEN','')
    frm   = os.environ.get('TWILIO_FROM_NUMBER','')
    if sid and token and frm:
        try:
            from twilio.rest import Client
            Client(sid, token).messages.create(
                body=f'Your Ethical Byte OTP is {otp}. Valid for {OTP_EXPIRY_MIN} minutes.',
                from_=frm, to=mobile
            )
        except Exception as e:
            print(f'[Twilio] {e}')
    print(f'[OTP] {mobile} → {otp}')
    return jsonify({'ok': True})

@app.route('/api/auth/verify-otp', methods=['POST'])
def verify_otp():
    data  = request.json or {}
    mobile = data.get('mobile','').strip()
    otp    = data.get('otp','').strip()
    with get_db() as db:
        row = db.execute(
            'SELECT * FROM otps WHERE mobile=? AND otp=? ORDER BY id DESC LIMIT 1',
            (mobile, otp)
        ).fetchone()
        if not row:
            return jsonify({'error': 'Invalid OTP'}), 400
        if datetime.utcnow().isoformat() > row['expires_at']:
            return jsonify({'error': 'OTP expired'}), 400
        existing = db.execute('SELECT id FROM users WHERE email=?', (row['email'],)).fetchone()
        if existing:
            return jsonify({'error': 'Email already registered'}), 409
        db.execute(
            'INSERT INTO users(name,email,password_hash,mobile) VALUES(?,?,?,?)',
            (row['name'], row['email'], row['password_hash'], mobile)
        )
        db.execute('DELETE FROM otps WHERE mobile=?', (mobile,))
    return jsonify({'ok': True})

@app.route('/api/auth/login', methods=['POST'])
def do_login():
    data  = request.json or {}
    email = data.get('email','').strip().lower()
    pw    = data.get('password','')
    with get_db() as db:
        u = db.execute('SELECT * FROM users WHERE email=?', (email,)).fetchone()
    if not u or u['password_hash'] != hash_pw(pw):
        return jsonify({'error': 'Invalid email or password'}), 401
    session['user_id'] = u['id']
    session['is_superuser'] = bool(u['is_superuser'])
    return jsonify({'ok': True, 'is_superuser': bool(u['is_superuser'])})

@app.route('/api/auth/logout', methods=['GET','POST'])
@app.route('/auth/logout')
def do_logout():
    session.clear()
    return redirect('/login')

@app.route('/api/auth/me')
@app.route('/auth/me')
def auth_me():
    if 'user_id' not in session:
        return jsonify({'authenticated': False}), 401
    with get_db() as db:
        u = db.execute('SELECT id,name,email,is_superuser FROM users WHERE id=?', (session['user_id'],)).fetchone()
    if not u:
        return jsonify({'authenticated': False}), 401
    return jsonify({'authenticated': True, 'user': dict(u)})

# ══════════════════════════════════════════════════════════════════════════
#  REPORT GENERATION
# ══════════════════════════════════════════════════════════════════════════

@app.route('/generate', methods=['POST'])
@login_required
def generate():
    data = request.json or {}
    fmt  = data.get('format', 'docx').lower()
    org  = data.get('orgName', 'Unknown Organisation')
    date = data.get('reportDate', datetime.today().strftime('%Y-%m-%d'))
    preparer = data.get('preparerName', '')
    exec_summary = data.get('execSummary', '')
    scope        = data.get('scope', '')
    vulns        = data.get('vulnerabilities', [])

    safe_org = re.sub(r'[^\w\s-]', '', org).strip().replace(' ', '_')
    filename = f'VAPT_Report_{safe_org}_{datetime.now().strftime("%Y%m%d%H%M%S")}.{fmt}'
    filepath = os.path.join(REPORT_DIR, filename)

    if fmt == 'docx':
        buf = generate_docx(org, date, preparer, exec_summary, scope, vulns)
    else:
        buf = generate_pdf(org, date, preparer, exec_summary, scope, vulns)

    with open(filepath, 'wb') as f:
        f.write(buf.getvalue())

    with get_db() as db:
        db.execute(
            'INSERT INTO reports(user_id,org_name,format,filename,filepath) VALUES(?,?,?,?,?)',
            (session['user_id'], org, fmt, filename, filepath)
        )

    buf.seek(0)
    mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if fmt == 'docx' else 'application/pdf'
    return send_file(buf, mimetype=mime, as_attachment=True, download_name=filename)

# ══════════════════════════════════════════════════════════════════════════
#  DOCX GENERATION
# ══════════════════════════════════════════════════════════════════════════


def get_logo_for_docx():
    """Convert logo to white background for DOCX, redrawing Ethical text which was black-on-black."""
    if not os.path.exists(LOGO_PATH):
        return None
    from PIL import Image as PILImg, ImageDraw, ImageFont
    import numpy as np
    img = PILImg.open(LOGO_PATH).convert('RGBA')
    arr = np.array(img)
    # Remove black background (make transparent)
    r, g, b = arr[:,:,0].astype(int), arr[:,:,1].astype(int), arr[:,:,2].astype(int)
    is_bg = (r < 20) & (g < 20) & (b < 20)
    arr[is_bg, 3] = 0
    # Composite onto white
    white = PILImg.new('RGBA', img.size, (255, 255, 255, 255))
    logo_t = PILImg.fromarray(arr)
    white.paste(logo_t, mask=logo_t)
    result = white.convert('RGB')
    # Redraw 'Ethical' text — it was black-on-black so it disappeared
    draw = ImageDraw.Draw(result)
    font_paths = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
    ]
    font = None
    for fp in font_paths:
        if os.path.exists(fp):
            font = ImageFont.truetype(fp, 26)
            break
    if font:
        bbox = draw.textbbox((0, 0), 'Ethical', font=font)
        text_h = bbox[3] - bbox[1]
        y = 12 + (38 - text_h) // 2
        draw.text((56, y), 'Ethical', fill=(15, 15, 15), font=font)
    buf = io.BytesIO()
    result.save(buf, format='PNG')
    buf.seek(0)
    return buf

def set_cell_shading(cell, fill_hex):
    """Apply background colour to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd if any
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex.lstrip('#'))
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def generate_docx(org, date, preparer, exec_summary, scope, vulns):
    """Generate DOCX from scratch, using the template only for header/footer/styles."""
    # ── Start from template to inherit header, footer, styles, page borders ──
    if os.path.exists(TEMPLATE_PATH):
        doc = DocxDocument(TEMPLATE_PATH)
        body = doc.element.body
        # Remove every child except the final sectPr (which holds header/footer/borders)
        sect_pr = body.find(qn('w:sectPr'))
        for child in list(body):
            if child is not sect_pr:
                body.remove(child)
    else:
        doc = DocxDocument()

    # ── Helpers ──────────────────────────────────────────────────────────────
    DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
    ACCENT    = RGBColor(0x44, 0x72, 0xC4)

    def add_para(text='', style='Normal', bold=False, size=None,
                 color=None, align=None, space_before=None, space_after=None,
                 italic=False):
        p = doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        pf = p.paragraph_format
        if space_before is not None: pf.space_before = Pt(space_before)
        if space_after  is not None: pf.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            if size:  run.font.size = Pt(size)
            if color: run.font.color.rgb = color
        return p

    def add_heading(text, level=1):
        return doc.add_heading(text, level=level)

    def page_break():
        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  COVER PAGE  (no header/footer — matches template notFirstPage)
    # ══════════════════════════════════════════════════════════════════════
    # Logo centred
    logo_buf = get_logo_for_docx()
    if logo_buf:
        doc.add_picture(logo_buf, width=Inches(2.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spacing
    for _ in range(3):
        add_para()

    # Main title
    for line in ('VULNERABILITY ASSESSMENT &', 'PENETRATION TESTING REPORT'):
        p = add_para(line, bold=True, size=22,
                     align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE,
                     space_before=0, space_after=4)

    add_para(space_before=12, space_after=0)

    # Info table
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ('Organisation',   org),
        ('Prepared By',    preparer),
        ('Report Date',    date),
        ('Classification', 'Confidential'),
    ]
    col_widths = [Inches(2.0), Inches(4.5)]
    for i, (k, v) in enumerate(rows_data):
        c0, c1 = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        c0.width = col_widths[0]
        c1.width = col_widths[1]
        # Label cell — dark blue bg, white text
        c0.paragraphs[0].clear()
        run0 = c0.paragraphs[0].add_run(k)
        run0.bold = True
        run0.font.size = Pt(11)
        run0.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(c0, '1F3864')
        set_cell_margins(c0)
        # Value cell
        c1.paragraphs[0].clear()
        run1 = c1.paragraphs[0].add_run(v)
        run1.font.size = Pt(11)
        set_cell_margins(c1)

    add_para()
    # Confidentiality note
    add_para(
        'This document is confidential and intended solely for the organisation named above. '
        'Unauthorised distribution is strictly prohibited.',
        italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x66, 0x66, 0x66)
    )

    page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════
    add_heading('Table of Contents', 1)
    toc_entries = [
        ('Executive Summary',        '2'),
        ('1. Vulnerability Summary',  '3'),
        ('2. Severity Distribution',  '4'),
        ('3. Vulnerability Details',  '5'),
    ]
    for entry, pg in toc_entries:
        dots = '.' * max(4, 55 - len(entry))
        p = add_para(size=10, space_before=4, space_after=4)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
        run = p.add_run(f'{entry} {dots} {pg}')
        run.font.size = Pt(10)

    page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    add_heading('Executive Summary', 1)
    add_para(exec_summary or 'No executive summary provided.', space_after=8)
    add_para()
    add_heading('Scope', 2)
    add_para(scope or 'No scope defined.', space_after=8)

    page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  1. VULNERABILITY SUMMARY TABLE
    # ══════════════════════════════════════════════════════════════════════
    add_heading('1. Vulnerability Summary', 1)

    if vulns:
        headers   = ['#', 'Title', 'Severity', 'OWASP', 'URLs', 'Status']
        col_w_in  = [0.4, 2.5, 0.9, 1.6, 0.9, 0.7]   # inches, total ≈ 7
        tbl = doc.add_table(rows=1 + len(vulns), cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, (h, w) in enumerate(zip(headers, col_w_in)):
            cell = tbl.rows[0].cells[i]
            cell.width = Inches(w)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '1F3864')
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)

        SEV_HEX = {'Critical':'C00000','High':'FF4444','Medium':'FFC000','Low':'92D050'}
        SEV_TXT = {'Critical': RGBColor(255,255,255), 'High': RGBColor(255,255,255),
                   'Medium': RGBColor(0,0,0), 'Low': RGBColor(0,0,0)}

        for i, v in enumerate(vulns):
            row  = tbl.rows[i + 1]
            sev  = v.get('severity', '')
            vals = [str(i+1), v.get('title',''), sev,
                    v.get('owasp',''), (v.get('urls','') or '')[:40], 'Open']
            for j, (val, w) in enumerate(zip(vals, col_w_in)):
                cell = row.cells[j]
                cell.width = Inches(w)
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                cell.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2, 5) else WD_ALIGN_PARAGRAPH.LEFT
                )
                set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
                if j == 2 and sev in SEV_HEX:
                    run.bold = True
                    run.font.color.rgb = SEV_TXT[sev]
                    set_cell_shading(cell, SEV_HEX[sev])
                elif i % 2 == 0:
                    set_cell_shading(cell, 'EEF2FA')
                else:
                    set_cell_shading(cell, 'F5F8FF')
    else:
        add_para('No vulnerabilities recorded.')

    page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  2. SEVERITY DISTRIBUTION
    # ══════════════════════════════════════════════════════════════════════
    add_heading('2. Severity Distribution', 1)

    chart_buf = make_bar_chart(vulns)
    if chart_buf:
        doc.add_picture(chart_buf, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para()

    # Counts summary table
    counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
    for v in vulns:
        s = v.get('severity', '')
        if s in counts: counts[s] += 1

    sum_tbl = doc.add_table(rows=2, cols=4)
    sum_tbl.style = 'Table Grid'
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    SEV_HEX2 = {'Critical':'C00000','High':'FF4444','Medium':'FFC000','Low':'92D050'}
    SEV_TXT2 = {'Critical': RGBColor(255,255,255), 'High': RGBColor(255,255,255),
                'Medium': RGBColor(0,0,0), 'Low': RGBColor(0,0,0)}
    for j, (sev, cnt) in enumerate(counts.items()):
        # Header cell
        hc = sum_tbl.rows[0].cells[j]
        hc.paragraphs[0].clear()
        hr = hc.paragraphs[0].add_run(sev)
        hr.bold = True; hr.font.size = Pt(11)
        hr.font.color.rgb = SEV_TXT2[sev]
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hc, SEV_HEX2[sev])
        set_cell_margins(hc, top=80, bottom=80)
        # Count cell
        vc = sum_tbl.rows[1].cells[j]
        vc.paragraphs[0].clear()
        vr = vc.paragraphs[0].add_run(str(cnt))
        vr.bold = True; vr.font.size = Pt(14)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(vc, top=80, bottom=80)

    page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  3. VULNERABILITY DETAILS
    # ══════════════════════════════════════════════════════════════════════
    add_heading('3. Vulnerability Details', 1)

    SEV_HEX3 = {'Critical':'C00000','High':'FF4444','Medium':'FFC000','Low':'92D050'}
    SEV_TXT3 = {'Critical': RGBColor(255,255,255), 'High': RGBColor(255,255,255),
                'Medium': RGBColor(0,0,0), 'Low': RGBColor(0,0,0)}

    for idx, v in enumerate(vulns, 1):
        sev   = v.get('severity', '')
        title = v.get('title', 'Untitled')

        # Title banner table (vuln name | severity badge)
        banner = doc.add_table(rows=1, cols=2)
        banner.style = 'Table Grid'
        banner.alignment = WD_TABLE_ALIGNMENT.CENTER
        bc0, bc1 = banner.rows[0].cells[0], banner.rows[0].cells[1]
        bc0.width = Inches(5.5)
        bc1.width = Inches(1.5)
        bc0.paragraphs[0].clear()
        r0 = bc0.paragraphs[0].add_run(f'{idx}. {title}')
        r0.bold = True; r0.font.size = Pt(12)
        r0.font.color.rgb = DARK_BLUE
        set_cell_shading(bc0, 'EEF2FA')
        set_cell_margins(bc0)
        bc1.paragraphs[0].clear()
        r1 = bc1.paragraphs[0].add_run(sev)
        r1.bold = True; r1.font.size = Pt(11)
        r1.font.color.rgb = SEV_TXT3.get(sev, RGBColor(0,0,0))
        bc1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(bc1, SEV_HEX3.get(sev, 'FFFFFF'))
        set_cell_margins(bc1)

        add_para(space_before=4, space_after=0)

        # Detail rows table
        detail_rows = [
            ('OWASP Mapping',  v.get('owasp', '')),
            ('Summary',        v.get('summary', '')),
            ('Risk',           v.get('risk', '')),
            ('Impact',         v.get('impact', '')),
            ('Advisory',       v.get('advisory', '')),
            ('Vulnerable URLs', v.get('urls', '')),
            ('Remediation',    v.get('remediation', '')),
        ]
        # Filter empty rows
        detail_rows = [(k, val) for k, val in detail_rows if str(val).strip()]

        if detail_rows:
            dtbl = doc.add_table(rows=len(detail_rows), cols=2)
            dtbl.style = 'Table Grid'
            dtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r, (k, val) in enumerate(detail_rows):
                c0 = dtbl.rows[r].cells[0]
                c1 = dtbl.rows[r].cells[1]
                c0.width = Inches(1.5)
                c1.width = Inches(5.5)
                c0.paragraphs[0].clear()
                lrun = c0.paragraphs[0].add_run(k)
                lrun.bold = True; lrun.font.size = Pt(10)
                set_cell_shading(c0, 'EEF2FA')
                set_cell_margins(c0)
                c1.paragraphs[0].clear()
                vrun = c1.paragraphs[0].add_run(str(val))
                vrun.font.size = Pt(10)
                set_cell_margins(c1)
                if r % 2 == 0:
                    set_cell_shading(c1, 'FAFCFF')
                else:
                    set_cell_shading(c1, 'F2F6FC')

        # POC screenshots
        poc_images = v.get('pocImages', [])
        if poc_images:
            add_para('Proof of Concept (Screenshots):', bold=True,
                     space_before=8, space_after=4)
            for img_b64 in poc_images:
                try:
                    img_data = base64.b64decode(img_b64.split(',')[-1])
                    img_buf  = io.BytesIO(img_data)
                    doc.add_picture(img_buf, width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

        if idx < len(vulns):
            page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════════
#  PDF GENERATION  — pixel-faithful to DOCX template
# ══════════════════════════════════════════════════════════════════════════

# ── Page constants ─────────────────────────────────────────────────────────
PAGE_W, PAGE_H = A4                      # 595.28 × 841.89 pt
MARGIN         = 72.0                    # 1 inch all sides
CONTENT_W      = PAGE_W - 2 * MARGIN
CONTENT_H      = PAGE_H - 2 * MARGIN
HEADER_H       = 35.4                    # pt from page top
# DOCX border: val=double sz=4(0.5pt) space=24(3pt from edge) offsetFrom=page
BORDER_SPACE   = 3.0                     # pt from page edge to outer line
BORDER_GAP     = 3.0                     # pt between outer and inner line
BORDER_LINE_W  = 0.5                     # pt per line
FOOTER_H       = 30.0                   # pt

# Colours
C_DARK   = colors.HexColor('#1F3864')
C_ACCENT = colors.HexColor('#4472C4')
C_WHITE  = colors.white
C_BLACK  = colors.black

SEV_RL = {
    'Critical': colors.HexColor('#C00000'),
    'High':     colors.HexColor('#FF4444'),
    'Medium':   colors.HexColor('#FFC000'),
    'Low':      colors.HexColor('#92D050'),
}

# ── Styles ──────────────────────────────────────────────────────────────────
def build_styles():
    base = getSampleStyleSheet()
    styles = {}

    def S(name, parent='Normal', **kw):
        styles[name] = ParagraphStyle(name, parent=base[parent], **kw)

    S('CoverTitle',  fontSize=22, leading=28, textColor=C_DARK,
      alignment=TA_CENTER, spaceAfter=6, fontName='Helvetica-Bold')
    S('CoverSub',    fontSize=13, leading=18, textColor=C_DARK,
      alignment=TA_CENTER, spaceAfter=4, fontName='Helvetica-Bold')
    S('CoverLabel',  fontSize=11, leading=16, textColor=C_BLACK,
      fontName='Helvetica-Bold')
    S('CoverValue',  fontSize=11, leading=16, textColor=C_BLACK,
      fontName='Helvetica')
    S('H1', 'Heading1', fontSize=16, leading=22, textColor=C_DARK,
      fontName='Helvetica-Bold', spaceBefore=14, spaceAfter=8,
      borderPad=0)
    S('H2', 'Heading2', fontSize=13, leading=18, textColor=C_ACCENT,
      fontName='Helvetica-Bold', spaceBefore=10, spaceAfter=6)
    S('Body', fontSize=10, leading=14, fontName='Helvetica',
      spaceAfter=6, alignment=TA_JUSTIFY)
    S('TH',  fontSize=9,  leading=12, fontName='Helvetica-Bold',
      textColor=C_WHITE, alignment=TA_CENTER)
    S('TD',  fontSize=9,  leading=12, fontName='Helvetica',
      alignment=TA_LEFT)
    S('TDC', fontSize=9,  leading=12, fontName='Helvetica',
      alignment=TA_CENTER)
    S('Label', fontSize=10, leading=14, fontName='Helvetica-Bold')
    S('Footer', fontSize=8, leading=10, fontName='Helvetica',
      alignment=TA_CENTER, textColor=colors.HexColor('#555555'))
    return styles


def draw_border(canvas, doc):
    """Double border on all pages except the cover (page 1)."""
    if canvas.getPageNumber() == 1:
        return
    canvas.saveState()
    canvas.setStrokeColor(C_BLACK)
    canvas.setLineWidth(BORDER_LINE_W)
    # DOCX double border: offsetFrom=page, space=3pt, gap=3pt between lines
    # Outer rect — 3pt from page edge
    o = BORDER_SPACE
    canvas.rect(o, o, PAGE_W - 2*o, PAGE_H - 2*o)
    # Inner rect — gap further in
    i = BORDER_SPACE + BORDER_GAP
    canvas.rect(i, i, PAGE_W - 2*i, PAGE_H - 2*i)
    canvas.restoreState()


def draw_header(canvas, doc):
    """Logo in header — all pages except cover."""
    if canvas.getPageNumber() == 1:
        return
    if not os.path.exists(LOGO_PATH):
        return
    canvas.saveState()
    logo_w = 87.75
    logo_h = 22.50
    x = PAGE_W - MARGIN - logo_w
    y = PAGE_H - HEADER_H + (HEADER_H - logo_h) / 2
    canvas.drawImage(LOGO_PATH, x, y, width=logo_w, height=logo_h,
                     preserveAspectRatio=True, mask='auto')
    # Thin separator line under header
    canvas.setStrokeColor(colors.HexColor('#CCCCCC'))
    canvas.setLineWidth(0.5)
    canvas.line(MARGIN, PAGE_H - HEADER_H, PAGE_W - MARGIN, PAGE_H - HEADER_H)
    canvas.restoreState()


def draw_footer(canvas, doc):
    """Centred 'Ethical Byte | N' footer — all pages except cover."""
    if canvas.getPageNumber() == 1:
        return
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.HexColor('#555555'))
    text = f'Ethical Byte  |  {canvas.getPageNumber()}'
    canvas.drawCentredString(PAGE_W / 2, 18, text)
    # Thin line above footer
    canvas.setStrokeColor(colors.HexColor('#CCCCCC'))
    canvas.setLineWidth(0.5)
    canvas.line(MARGIN, 28, PAGE_W - MARGIN, 28)
    canvas.restoreState()


def on_page(canvas, doc):
    draw_border(canvas, doc)
    draw_header(canvas, doc)
    draw_footer(canvas, doc)


def make_bar_chart(vulns):
    """3-D style bar chart of severity distribution, returns BytesIO PNG."""
    counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
    for v in vulns:
        s = v.get('severity', '')
        if s in counts:
            counts[s] += 1

    labels = list(counts.keys())
    values = [counts[k] for k in labels]
    bar_colors = [SEV_COLORS[k] for k in labels]

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(labels, values, color=bar_colors, edgecolor='white', linewidth=0.8, zorder=3, width=0.5)
    # Pseudo-3D shadow
    for bar in bars:
        shadow = mpatches.FancyBboxPatch(
            (bar.get_x() + 0.04, 0), bar.get_width(), bar.get_height(),
            boxstyle='square,pad=0', color='#00000022', zorder=2
        )
        ax.add_patch(shadow)
    ax.set_facecolor('#F5F8FF')
    fig.patch.set_facecolor('white')
    ax.set_ylabel('Count', fontsize=11)
    ax.set_title('Severity Distribution', fontsize=13, fontweight='bold', color='#1F3864')
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    ax.grid(axis='y', linestyle='--', alpha=0.5, zorder=0)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    for bar, val in zip(bars, values):
        if val:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                    str(val), ha='center', va='bottom', fontweight='bold', fontsize=10)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_pdf(org, date, preparer, exec_summary, scope, vulns):
    S = build_styles()
    buf = io.BytesIO()

    # top/bottom margins account for header & footer bands
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=MARGIN + HEADER_H,
        bottomMargin=FOOTER_H + MARGIN * 0.5,
        title=f'VAPT Report — {org}',
        author='Ethical Byte',
    )

    story = []

    # ── COVER PAGE ────────────────────────────────────────────────────────
    story.append(Spacer(1, 30))
    if os.path.exists(LOGO_PATH):
        img = Image(LOGO_PATH, width=2.5*inch, height=0.65*inch)
        img.hAlign = 'CENTER'
        story.append(img)
    story.append(Spacer(1, 30))
    story.append(Paragraph('VULNERABILITY ASSESSMENT &amp;', S['CoverTitle']))
    story.append(Paragraph('PENETRATION TESTING REPORT', S['CoverTitle']))
    story.append(Spacer(1, 30))

    # Info table
    cover_data = [
        [Paragraph('Organisation',    S['CoverLabel']), Paragraph(org,           S['CoverValue'])],
        [Paragraph('Prepared By',     S['CoverLabel']), Paragraph(preparer,      S['CoverValue'])],
        [Paragraph('Report Date',     S['CoverLabel']), Paragraph(date,          S['CoverValue'])],
        [Paragraph('Classification',  S['CoverLabel']), Paragraph('Confidential',S['CoverValue'])],
    ]
    cover_tbl = Table(cover_data, colWidths=[CONTENT_W*0.35, CONTENT_W*0.65])
    cover_tbl.setStyle(TableStyle([
        ('BACKGROUND',  (0,0),(-1,-1), colors.HexColor('#F0F4FA')),
        ('BACKGROUND',  (0,0),(0,-1),  colors.HexColor('#1F3864')),
        ('TEXTCOLOR',   (0,0),(0,-1),  C_WHITE),
        ('GRID',        (0,0),(-1,-1), 0.5, colors.HexColor('#C0C8D8')),
        ('ROWBACKGROUNDS', (0,0),(-1,-1), [colors.HexColor('#EEF2FA'), colors.HexColor('#F5F8FF')]),
        ('VALIGN',      (0,0),(-1,-1), 'MIDDLE'),
        ('LEFTPADDING', (0,0),(-1,-1), 10),
        ('RIGHTPADDING',(0,0),(-1,-1), 10),
        ('TOPPADDING',  (0,0),(-1,-1), 8),
        ('BOTTOMPADDING',(0,0),(-1,-1), 8),
    ]))
    story.append(cover_tbl)
    story.append(Spacer(1, 20))
    # Confidentiality notice
    story.append(Paragraph(
        '<i>This document is confidential and intended solely for the organisation named above. '
        'Unauthorised distribution is strictly prohibited.</i>',
        ParagraphStyle('Note', parent=S['Body'], fontSize=9, textColor=colors.HexColor('#666666'),
                       alignment=TA_CENTER)
    ))
    story.append(PageBreak())

    # ── TABLE OF CONTENTS (manual) ────────────────────────────────────────
    story.append(Paragraph('Table of Contents', S['H1']))
    toc_data = [
        ['Executive Summary', '2'],
        ['1. Vulnerability Summary', '3'],
        ['2. Severity Distribution', '4'],
        ['3. Vulnerability Details', '5'],
    ]
    for entry, pg in toc_data:
        story.append(Paragraph(
            f'{entry} {"." * max(1, 60-len(entry))} {pg}',
            ParagraphStyle('TOC', parent=S['Body'], fontName='Helvetica',
                           fontSize=10, leading=16)
        ))
    story.append(PageBreak())

    # ── EXECUTIVE SUMMARY ─────────────────────────────────────────────────
    story.append(Paragraph('Executive Summary', S['H1']))
    story.append(HRFlowable(width=CONTENT_W, thickness=1, color=C_DARK, spaceAfter=8))
    story.append(Paragraph(exec_summary or 'No executive summary provided.', S['Body']))
    story.append(Spacer(1, 12))
    story.append(Paragraph('Scope', S['H2']))
    story.append(Paragraph(scope or 'No scope defined.', S['Body']))
    story.append(PageBreak())

    # ── VULNERABILITY SUMMARY TABLE ───────────────────────────────────────
    story.append(Paragraph('1. Vulnerability Summary', S['H1']))
    story.append(HRFlowable(width=CONTENT_W, thickness=1, color=C_DARK, spaceAfter=8))

    if vulns:
        col_w = [CONTENT_W*x for x in [0.06, 0.32, 0.13, 0.28, 0.12, 0.09]]
        hdr = [Paragraph(h, S['TH']) for h in ['#','Title','Severity','OWASP','URLs','Status']]
        rows = [hdr]
        ts = [
            ('BACKGROUND', (0,0),(-1,0), C_DARK),
            ('TEXTCOLOR',  (0,0),(-1,0), C_WHITE),
            ('GRID',       (0,0),(-1,-1), 0.4, colors.HexColor('#C0C8D8')),
            ('VALIGN',     (0,0),(-1,-1), 'MIDDLE'),
            ('LEFTPADDING',(0,0),(-1,-1), 6),
            ('RIGHTPADDING',(0,0),(-1,-1), 6),
            ('TOPPADDING', (0,0),(-1,-1), 5),
            ('BOTTOMPADDING',(0,0),(-1,-1), 5),
            ('ROWBACKGROUNDS', (0,1),(-1,-1),
             [colors.HexColor('#F5F8FF'), colors.HexColor('#EEF2FA')]),
        ]
        for i, v in enumerate(vulns):
            sev = v.get('severity','')
            sev_color = SEV_RL.get(sev, colors.white)
            txt_col   = C_WHITE if sev in ('Critical','High') else C_BLACK
            row = [
                Paragraph(str(i+1), S['TDC']),
                Paragraph(v.get('title',''), S['TD']),
                Paragraph(f'<b>{sev}</b>', ParagraphStyle('SC', parent=S['TDC'],
                    textColor=txt_col, fontName='Helvetica-Bold', fontSize=9)),
                Paragraph(v.get('owasp',''), S['TD']),
                Paragraph((v.get('urls','') or '')[:40] + ('…' if len(v.get('urls',''))>40 else ''), S['TD']),
                Paragraph('Open', S['TDC']),
            ]
            rows.append(row)
            ts.append(('BACKGROUND', (2, i+1), (2, i+1), sev_color))

        tbl = Table(rows, colWidths=col_w, repeatRows=1)
        tbl.setStyle(TableStyle(ts))
        story.append(tbl)
    else:
        story.append(Paragraph('No vulnerabilities recorded.', S['Body']))
    story.append(PageBreak())

    # ── SEVERITY CHART ────────────────────────────────────────────────────
    story.append(Paragraph('2. Severity Distribution', S['H1']))
    story.append(HRFlowable(width=CONTENT_W, thickness=1, color=C_DARK, spaceAfter=8))
    chart_buf = make_bar_chart(vulns)
    if chart_buf:
        chart_img = Image(chart_buf, width=CONTENT_W * 0.85, height=CONTENT_W * 0.42)
        chart_img.hAlign = 'CENTER'
        story.append(chart_img)

    # Counts summary under chart
    counts = {'Critical':0,'High':0,'Medium':0,'Low':0}
    for v in vulns:
        s = v.get('severity','')
        if s in counts: counts[s] += 1
    summary_data = [[Paragraph(f'<b>{k}</b>', ParagraphStyle('SC2', parent=S['TDC'],
        textColor=C_WHITE if k in ('Critical','High') else C_BLACK,
        fontName='Helvetica-Bold', fontSize=10))
        for k in counts.keys()],
        [Paragraph(str(c), ParagraphStyle('SV', parent=S['TDC'],
            fontSize=12, fontName='Helvetica-Bold'))
        for c in counts.values()]]
    sum_tbl = Table(summary_data, colWidths=[CONTENT_W/4]*4)
    sum_ts = [
        ('ALIGN',  (0,0),(-1,-1), 'CENTER'),
        ('VALIGN', (0,0),(-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0),(-1,-1), 8),
        ('BOTTOMPADDING',(0,0),(-1,-1), 8),
        ('ROUNDEDCORNERS', [4,4,4,4]),
    ]
    for i, k in enumerate(['Critical','High','Medium','Low']):
        sum_ts.append(('BACKGROUND', (i,0),(i,0), SEV_RL[k]))
        sum_ts.append(('BACKGROUND', (i,1),(i,1), colors.HexColor('#F5F8FF')))
        sum_ts.append(('BOX', (i,0),(i,1), 0.5, colors.HexColor('#C0C8D8')))
    sum_tbl.setStyle(TableStyle(sum_ts))
    story.append(Spacer(1, 16))
    story.append(sum_tbl)
    story.append(PageBreak())

    # ── DETAILED FINDINGS ─────────────────────────────────────────────────
    story.append(Paragraph('3. Vulnerability Details', S['H1']))
    story.append(HRFlowable(width=CONTENT_W, thickness=1, color=C_DARK, spaceAfter=8))

    for idx, v in enumerate(vulns, 1):
        sev   = v.get('severity','')
        title = v.get('title','Untitled')
        sev_color = SEV_RL.get(sev, colors.HexColor('#AAAAAA'))
        txt_col   = C_WHITE if sev in ('Critical','High') else C_BLACK

        # Vuln title banner
        title_data = [[
            Paragraph(f'<b>{idx}. {title}</b>',
                ParagraphStyle('VTitle', parent=S['H2'], textColor=C_DARK,
                               fontSize=12, fontName='Helvetica-Bold')),
            Paragraph(f'<b>{sev}</b>',
                ParagraphStyle('VS', parent=S['TDC'], textColor=txt_col,
                               fontName='Helvetica-Bold', fontSize=10)),
        ]]
        title_tbl = Table(title_data, colWidths=[CONTENT_W*0.75, CONTENT_W*0.25])
        title_tbl.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (0,0), colors.HexColor('#EEF2FA')),
            ('BACKGROUND', (1,0), (1,0), sev_color),
            ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING',(0,0), (-1,-1), 10),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOTTOMPADDING',(0,0),(-1,-1), 8),
            ('BOX',        (0,0), (-1,-1), 0.5, colors.HexColor('#C0C8D8')),
        ]))
        story.append(title_tbl)
        story.append(Spacer(1, 6))

        # Detail rows
        fields = [
            ('OWASP Mapping', v.get('owasp','')),
            ('Summary',       v.get('summary','')),
            ('Risk',          v.get('risk','')),
            ('Impact',        v.get('impact','')),
            ('Advisory',      v.get('advisory','')),
            ('Vulnerable URLs',v.get('urls','')),
            ('Remediation',   v.get('remediation','')),
        ]
        detail_data = []
        for label, val in fields:
            if val:
                detail_data.append([
                    Paragraph(label, S['Label']),
                    Paragraph(str(val), S['Body']),
                ])
        if detail_data:
            dtbl = Table(detail_data, colWidths=[CONTENT_W*0.22, CONTENT_W*0.78])
            dtbl.setStyle(TableStyle([
                ('GRID',    (0,0),(-1,-1), 0.4, colors.HexColor('#D0D8E8')),
                ('BACKGROUND', (0,0),(0,-1), colors.HexColor('#EEF2FA')),
                ('VALIGN',  (0,0),(-1,-1), 'TOP'),
                ('LEFTPADDING',(0,0),(-1,-1), 8),
                ('RIGHTPADDING',(0,0),(-1,-1), 8),
                ('TOPPADDING',(0,0),(-1,-1), 6),
                ('BOTTOMPADDING',(0,0),(-1,-1), 6),
                ('ROWBACKGROUNDS', (1,0),(1,-1),
                 [colors.HexColor('#FAFCFF'), colors.HexColor('#F2F6FC')]),
            ]))
            story.append(dtbl)

        # POC screenshots
        poc_images = v.get('pocImages', [])
        if poc_images:
            story.append(Spacer(1, 8))
            story.append(Paragraph('Proof of Concept (Screenshots):', S['Label']))
            story.append(Spacer(1, 4))
            for img_b64 in poc_images:
                try:
                    img_data = base64.b64decode(img_b64.split(',')[-1])
                    img_buf  = io.BytesIO(img_data)
                    pil_img  = PILImage.open(img_buf)
                    w, h     = pil_img.size
                    max_w    = CONTENT_W
                    scale    = min(1.0, max_w / w)
                    disp_w   = w * scale
                    disp_h   = h * scale
                    img_buf.seek(0)
                    img_flow = Image(img_buf, width=disp_w, height=disp_h)
                    img_flow.hAlign = 'CENTER'
                    story.append(img_flow)
                    story.append(Spacer(1, 6))
                except Exception:
                    pass

        if idx < len(vulns):
            story.append(Spacer(1, 20))
            story.append(HRFlowable(width=CONTENT_W, thickness=0.5,
                                    color=colors.HexColor('#C0C8D8'), spaceAfter=12))

    # Build PDF
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════════
#  REPORTS API
# ══════════════════════════════════════════════════════════════════════════

@app.route('/api/reports')
@login_required
def list_reports():
    with get_db() as db:
        u = db.execute('SELECT is_superuser FROM users WHERE id=?', (session['user_id'],)).fetchone()
        if u and u['is_superuser']:
            rows = db.execute('SELECT r.*,u.name as user_name FROM reports r JOIN users u ON r.user_id=u.id ORDER BY r.created_at DESC').fetchall()
        else:
            rows = db.execute('SELECT * FROM reports WHERE user_id=? ORDER BY created_at DESC', (session['user_id'],)).fetchall()
    return jsonify([dict(r) for r in rows])

@app.route('/api/reports/<int:rid>', methods=['DELETE'])
@login_required
def delete_report(rid):
    with get_db() as db:
        u  = db.execute('SELECT is_superuser FROM users WHERE id=?', (session['user_id'],)).fetchone()
        r  = db.execute('SELECT * FROM reports WHERE id=?', (rid,)).fetchone()
        if not r:
            return jsonify({'error': 'Not found'}), 404
        if not (u['is_superuser'] or r['user_id'] == session['user_id']):
            return jsonify({'error': 'Forbidden'}), 403
        try:
            os.remove(r['filepath'])
        except Exception:
            pass
        db.execute('DELETE FROM reports WHERE id=?', (rid,))
    return jsonify({'ok': True})

@app.route('/api/reports/<int:rid>/download')
@login_required
def download_report(rid):
    with get_db() as db:
        u = db.execute('SELECT is_superuser FROM users WHERE id=?', (session['user_id'],)).fetchone()
        r = db.execute('SELECT * FROM reports WHERE id=?', (rid,)).fetchone()
    if not r:
        return jsonify({'error': 'Not found'}), 404
    if not (u['is_superuser'] or r['user_id'] == session['user_id']):
        return jsonify({'error': 'Forbidden'}), 403
    return send_file(r['filepath'], as_attachment=True, download_name=r['filename'])

@app.route('/api/stats')
@login_required
def stats():
    with get_db() as db:
        u = db.execute('SELECT is_superuser FROM users WHERE id=?', (session['user_id'],)).fetchone()
        if u and u['is_superuser']:
            total  = db.execute('SELECT COUNT(*) FROM reports').fetchone()[0]
            docx_c = db.execute("SELECT COUNT(*) FROM reports WHERE format='docx'").fetchone()[0]
            pdf_c  = db.execute("SELECT COUNT(*) FROM reports WHERE format='pdf'").fetchone()[0]
        else:
            uid    = session['user_id']
            total  = db.execute('SELECT COUNT(*) FROM reports WHERE user_id=?',(uid,)).fetchone()[0]
            docx_c = db.execute("SELECT COUNT(*) FROM reports WHERE user_id=? AND format='docx'",(uid,)).fetchone()[0]
            pdf_c  = db.execute("SELECT COUNT(*) FROM reports WHERE user_id=? AND format='pdf'",(uid,)).fetchone()[0]
    return jsonify({'total': total, 'docx': docx_c, 'pdf': pdf_c})

# ══════════════════════════════════════════════════════════════════════════
#  USERS API  (superuser only)
# ══════════════════════════════════════════════════════════════════════════

@app.route('/api/users')
@superuser_required
def list_users():
    with get_db() as db:
        rows = db.execute('SELECT id,name,email,mobile,is_superuser,created_at FROM users ORDER BY id').fetchall()
    return jsonify([dict(r) for r in rows])

@app.route('/api/users', methods=['POST'])
@superuser_required
def create_user():
    data = request.json or {}
    name  = data.get('name','').strip()
    email = data.get('email','').strip().lower()
    pw    = data.get('password','')
    if not all([name, email, pw]):
        return jsonify({'error': 'name, email and password required'}), 400
    with get_db() as db:
        try:
            db.execute('INSERT INTO users(name,email,password_hash) VALUES(?,?,?)',
                       (name, email, hash_pw(pw)))
        except sqlite3.IntegrityError:
            return jsonify({'error': 'Email already exists'}), 409
    return jsonify({'ok': True})

@app.route('/api/users/<int:uid>', methods=['DELETE'])
@superuser_required
def delete_user(uid):
    if uid == session['user_id']:
        return jsonify({'error': 'Cannot delete yourself'}), 400
    with get_db() as db:
        db.execute('DELETE FROM users WHERE id=?', (uid,))
    return jsonify({'ok': True})

@app.route('/api/users/<int:uid>/reset-password', methods=['POST'])
@superuser_required
def reset_password(uid):
    data = request.json or {}
    pw   = data.get('password','')
    if not pw:
        return jsonify({'error': 'password required'}), 400
    with get_db() as db:
        db.execute('UPDATE users SET password_hash=? WHERE id=?', (hash_pw(pw), uid))
    return jsonify({'ok': True})

# ══════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
